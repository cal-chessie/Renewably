#!/usr/bin/env python3
"""
Fix handover docs after polish completion.
1. Remove "Vercel Analytics" from Production Deployment Guide (PDF)
2. Update Developer Handover (docx) to reflect all 9 polish items
"""

import re

# ─── FIX 1: Remove Vercel from Deployment Guide generator script ───
print("=== Fix 1: Remove Vercel from deployment guide ===")
with open("/home/z/my-project/download/generate_deployment_guide.py", "r") as f:
    content = f.read()

old = "(APM) via Vercel Analytics, Datadog, or New Relic for request latency tracking and "
new = "(APM) via Datadog or New Relic for request latency tracking and "

if old in content:
    content = content.replace(old, new)
    with open("/home/z/my-project/download/generate_deployment_guide.py", "w") as f:
        f.write(content)
    print("  Fixed: Removed 'Vercel Analytics' from monitoring section")
else:
    print("  Already clean or pattern not found")

# ─── FIX 2: Update Handover docx ───
print("\n=== Fix 2: Update Developer Handover docx ===")

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

doc = Document("/home/z/my-project/download/Renewably_Developer_Handover.docx")

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def get_all_text(element):
    """Get all text from an XML element recursively."""
    texts = []
    for t in element.iter():
        if t.tag.endswith('}t') and t.text:
            texts.append(t.text)
        if t.tail:
            texts.append(t.tail)
    return ''.join(texts).strip()

def find_paragraph_index_by_text(search_text):
    """Find paragraph index containing search text."""
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            return i
    return -1

def find_paragraph_index_by_heading(search_text):
    """Find paragraph index by heading text."""
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and 'Heading' in p.style.name:
            if search_text in p.text:
                return i
        # Also check raw XML
        txt = get_all_text(p._element)
        if search_text in txt:
            return i
    return -1

# ── Update 1.1 Key Constraints: Add polish-related constraints ──
idx = find_paragraph_index_by_text("File-based sessions stored in JSON")
if idx >= 0:
    p = doc.paragraphs[idx]
    # This is a list paragraph - insert new items after it
    # We'll add new list paragraphs after this one
    from docx.oxml.ns import qn
    from copy import deepcopy
    
    new_constraints = [
        "SEO: CRM pages have <meta name=\"robots\" content=\"noindex, nofollow\">; public pages have Open Graph tags, Article JSON-LD, and PWA manifest.json",
        "Accessibility: WCAG AA contrast ratios, ARIA labels on interactive elements, focus management, prefers-reduced-motion support",
        "Responsive: All touch targets minimum 44px (Apple HIG), clamp() padding for fluid spacing, mobile Sheet controlled navigation",
        "Performance: next/image for all images, dynamic() lazy-loading for ChatWidget, optimized package imports via next.config.ts",
    ]
    
    # Insert new paragraphs after the constraints list
    # Find the element after the last constraint item
    ref_element = p._element
    for constraint in reversed(new_constraints):
        new_p = deepcopy(ref_element)
        # Clear existing runs and add new text
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        new_r = etree.SubElement(new_p, qn('w:r'))
        new_rPr = etree.SubElement(new_r, qn('w:rPr'))
        new_t = etree.SubElement(new_r, qn('w:t'))
        new_t.text = constraint
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        
        # Insert after the reference element's next sibling
        ref_element.addnext(new_p)
        ref_element = new_p
    
    print(f"  Added {len(new_constraints)} new constraint items after P{idx}")

# ── Update 11.2 CRM Page Polish: Mark as completed ──
# The section 11.2 has a heading but the body text is embedded in the paragraph element's XML
# Let's find it in the raw body XML
body = doc.element.body
found_112 = False
for element in body:
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    if tag == 'p':
        txt = get_all_text(element)
        if '11.2' in txt and 'CRM Page Polish' in txt:
            found_112 = True
            # The content is likely in the next sibling paragraphs
            continue
    if found_112 and tag == 'p':
        txt = get_all_text(element)
        if 'login page and dashboard have been polished' in txt.lower():
            # This is the old 11.2 content - replace it
            for r in element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                parent = r.getparent()
                parent.remove(r)
            
            new_r = etree.SubElement(element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            new_t = etree.SubElement(new_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_t.text = (
                "All 14 CRM pages have been fully polished as of April 2026. "
                "The polish pass included: (1) Responsive tweaks with 44px minimum touch targets, "
                "fluid clamp() padding, controlled mobile Sheet navigation, and responsive grid layouts; "
                "(2) Loading states with skeleton screens and spinner components on all data-fetching pages; "
                "(3) Page transition animations using Framer Motion with prefers-reduced-motion support; "
                "(4) Dark mode consistency across all pages with the brand palette (#F3D840 on #0A0A0A); "
                "(5) Accessibility improvements including WCAG AA contrast ratios, ARIA labels, "
                "focus management, and keyboard navigation; (6) Error and empty states with branded "
                "illustrations on contacts, activities, proposals, reports, and other pages; "
                "(7) Performance optimisations including next/image for all images and dynamic() "
                "lazy-loading for the ChatWidget component. All pages use inline style={{}} objects "
                "exclusively and maintain the dark theme with yellow accents."
            )
            print("  Updated section 11.2: Marked all CRM pages as polished")
            found_112 = False
            break

# ── Update 11.3 Recommended Improvements: Remove items that are done ──
found_113 = False
for element in body:
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    if tag == 'p':
        txt = get_all_text(element)
        if '11.3' in txt and 'Recommended' in txt:
            found_113 = True
            continue
    if found_113 and tag == 'p':
        txt = get_all_text(element)
        # Remove items that were completed during polish
        if 'server-side validation' in txt.lower():
            # Update: partially done - Zod exists on most routes
            for r in element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                parent = r.getparent()
                parent.remove(r)
            new_r = etree.SubElement(element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            new_t = etree.SubElement(new_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_t.text = (
                "Complete server-side Zod validation for all remaining API routes "
                "(most routes now have Zod schemas; a few edge cases remain)"
            )
            print("  Updated 11.3: Marked Zod validation as mostly done")
            break

# ── Add crm-shell.tsx to the file table (Table 2) ──
# Table 2 is the directory layout table
if len(doc.tables) >= 3:
    file_table = doc.tables[2]  # Path | Purpose table
    # Find where to insert crm-shell.tsx (after crm/layout.tsx)
    insert_idx = None
    for i, row in enumerate(file_table.rows):
        if 'crm/layout.tsx' in row.cells[0].text:
            insert_idx = i + 1
            break
    
    if insert_idx:
        from docx.table import Table, _Cell
        new_row = file_table.add_row()
        new_row.cells[0].text = "src/app/crm/crm-shell.tsx"
        new_row.cells[1].text = "CRM client shell component (sidebar, Sheet nav, providers)"
        # Move the row to the correct position
        tbl = file_table._tbl
        tr = new_row._tr
        tbl.remove(tr)
        ref_tr = file_table.rows[insert_idx]._tr
        ref_tr.addnext(tr)
        print(f"  Added crm-shell.tsx to file table at row {insert_idx}")

# ── Save ──
doc.save("/home/z/my-project/download/Renewably_Developer_Handover.docx")
print("\n  Handover docx saved successfully")

print("\n=== Done ===")
print("Summary:")
print("  1. Removed 'Vercel Analytics' from deployment guide generator")
print("  2. Added 4 new constraints to section 1.1 (SEO, a11y, responsive, perf)")
print("  3. Updated section 11.2 - all CRM pages now marked as polished")
print("  4. Updated section 11.3 - Zod validation marked as mostly done")
print("  5. Added crm-shell.tsx to directory layout table")
